# -*- coding: utf-8 -*-
"""
Created on Sat Mar 21 14:46:22 2026

@author: hp
"""

import requests
from requests.auth import HTTPBasicAuth
import data_extraction_const as const
import pandas as pd
from bs4 import BeautifulSoup
import uuid
import os
import docx2txt
from langchain_core.documents import Document
import fitz  # PyMuPDF
import docx
from datetime import datetime
import pytz

def get_confluence_response(url):
    response = requests.get(
        url,
        auth=HTTPBasicAuth(const.CONFLUENCE_USERNAME, const.CONFLUENCE_KEY),
        headers={"Accept": "application/json"}
    )
    
    print(response.status_code)
    print(response.text)
    
    
def get_all_pages():
    url = const.CONFLUENCE_BASE_URL + const.CONFLUENCE_ALL_PAGES_API_ENDPOINT
    auth = HTTPBasicAuth(const.CONFLUENCE_USERNAME, const.CONFLUENCE_KEY)
    start = 0
    limit = 25
    all_pages = []

    while True:
        params = {
            "start": start,
            "limit": limit,
            "expand": "history,space,version,body.storage"
        }

        response = requests.get(
            url,
            auth=auth,
            headers={"Accept": "application/json"},
            params=params
        )

        if response.status_code != 200:
            print("Error:", response.status_code, response.text)
            break

        data = response.json()
        results = data.get("results", [])

        if not results:
            break

        for page in results:
            page_data = {
                "page": page.get("id"),
                "title": page.get("title"),
                "content": page.get("body", {}).get("storage", {}).get("value"),
                "author": page.get("history", {}).get("createdBy", {}).get("displayName"),
                "created_on": page.get("history", {}).get("createdDate"),
                "last_modified_on": page.get("version", {}).get("when"),
                "space": page.get("space", {}).get("name")
            }
            all_pages.append(page_data)

        start += limit
    
    df = pd.DataFrame(all_pages)

    df.to_csv("D:\\Coding\\GEN AI - Transformation\\Data-Extraction\\confluence_pages.csv", index=False)    


def clean_html(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    return soup.get_text(separator=" ").strip()


def chunk_text(text, chunk_size=300, overlap=50):
    words = text.split()
    chunks = []

    if overlap >= chunk_size:
        raise ValueError("overlap must be smaller than chunk_size")

    start = 0
    n = len(words)

    while start < n:
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)

        # move forward with overlap
        start += (chunk_size - overlap)

    return chunks


def get_all_chuks_pages_from_confluence():
    url = const.CONFLUENCE_BASE_URL + const.CONFLUENCE_ALL_PAGES_API_ENDPOINT
    auth = HTTPBasicAuth(const.CONFLUENCE_USERNAME, const.CONFLUENCE_KEY)

    start = 0
    limit = 25
    final_rows = []

    while True:
        params = {
            "start": start,
            "limit": limit,
            "expand": "history,space,version,body.storage,metadata.labels"
        }

        response = requests.get(
            url,
            auth=auth,
            headers={"Accept": "application/json"},
            params=params
        )

        if response.status_code != 200:
            print("Error:", response.status_code, response.text)
            break

        data = response.json()
        results = data.get("results", [])

        if not results:
            break

        for page in results:
            page_id = page.get("id")
            title = page.get("title")

            html_content = page.get("body", {}).get("storage", {}).get("value", "")
            clean_text_data = clean_html(html_content)

            chunks = chunk_text(clean_text_data,chunk_size=100, overlap=25)

            labels = [lbl.get("name") for lbl in page.get("metadata", {}).get("labels", {}).get("results", [])]

            for idx, chunk in enumerate(chunks):
                row = {
                    "page": page_id,
                    "title": title,
                    "content": html_content,
                    "clean_text": clean_text_data,
                    "chunk_id": str(uuid.uuid4()),
                    "chunk_index": idx,
                    "chunk_text": chunk,
                    "author": page.get("history", {}).get("createdBy", {}).get("displayName"),
                    "created_on": page.get("history", {}).get("createdDate"),
                    "last_modified_on": page.get("version", {}).get("when"),
                    "space": page.get("space", {}).get("name"),
                    "space_key": page.get("space", {}).get("key"),
                    "labels": ",".join(labels),
                    "url": f"{const.CONFLUENCE_BASE_URL}/pages/{page_id}",
                    "word_count": len(clean_text_data.split()),
                    "summary": ""  # placeholder for future LLM summary
                }
                final_rows.append(row)

        start += limit
    
    df = pd.DataFrame(final_rows)

    df.to_csv("D:\\Coding\\GEN AI - Transformation\\Data-Extraction\\confluence_pages_2.csv", index=False)
    return final_rows

def to_utc(local_time_str: str, tz_name: str):
    """
    Convert a local time string to UTC.
    local_time_str format: 'YYYY-MM-DD HH:MM'
    tz_name: e.g. 'Asia/Kolkata', 'America/New_York'
    """
    # Parse the local time
    local_tz = pytz.timezone(tz_name)
    naive_dt = datetime.strptime(local_time_str, "%Y-%m-%d %H:%M:%S")
    
    # Localize and convert to UTC
    local_dt = local_tz.localize(naive_dt)
    utc_dt = local_dt.astimezone(pytz.utc)
    
    return utc_dt

def extract_docx_content(doc_path: str):
    """
    Extract text and images from a Word document (.docx).
    Returns a LangChain Document with metadata including image paths.
    """
    
    image_dir= const.extracted_image
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"{doc_path} not found")

    # Ensure image directory exists before extraction
    os.makedirs(image_dir, exist_ok=True)

    stat_info = os.stat(doc_path)
    created_fs = to_utc(datetime.fromtimestamp(stat_info.st_ctime).strftime("%Y-%m-%d %H:%M:%S"),"Asia/Kolkata")
    modified_fs = datetime.fromtimestamp(stat_info.st_mtime)
    filesize= stat_info.st_size
    # Extract text and images (images saved into image_dir if any)
    text = docx2txt.process(doc_path, image_dir)
    
    doc_meta = docx.Document(doc_path)
    props = doc_meta.core_properties
    created_date =props.created
    author= props.author
    # Collect image paths only if directory has files
    images = []
    if os.path.exists(image_dir):
        images = [os.path.join(image_dir, f) for f in os.listdir(image_dir) if os.path.isfile(os.path.join(image_dir, f))]

    return Document(
        page_content=text,
        metadata={"source": doc_path, 
                  "images": images,
                  "created_date":str(created_fs) if created_fs else None,
                  "utc_created_date":str(created_date),
                  "last_modified_date":str(modified_fs) if modified_fs else None,
                  "filesize":filesize,
                  'author':author
                  }
    )

def extract_pdf_content(pdf_path: str):
    """
    Extract text and images from a PDF file.
    Returns a LangChain Document with metadata including image paths.
    """
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"{pdf_path} not found")

    doc = fitz.open(pdf_path)
    text = ""
    images = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text += page.get_text("text")

        # Extract images
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            pix = fitz.Pixmap(doc, xref)
            img_path = f"pdf_image_{page_num}_{img_index}.png"
            pix.save(img_path)
            images.append(img_path)

    return Document(
        page_content=text,
        metadata={"source": pdf_path, "images": images}
    )


import psycopg2
from pgvector.psycopg2 import register_vector
from sentence_transformers import SentenceTransformer

# Get this from Supabase → Settings → Database → Connection string
conn = psycopg2.connect(f"postgresql://postgres:{const.SUPERBASE_DB_PASSWORD}@db.RAGBasedLearning.supabase.co:5432/postgres")
register_vector(conn)

model = SentenceTransformer("BAAI/bge-base-en-v1.5")  # 768 dims, runs on CPU fine

def upsert_chunk(doc_id, source, title, url, content):
    vector = model.encode(content, normalize_embeddings=True).tolist()
    with conn.cursor() as cur:
        cur.execute("""
            INSERT INTO rag_chunks (doc_id, source, title, url, content, embedding)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON CONFLICT (doc_id) DO UPDATE
            SET content=EXCLUDED.content, embedding=EXCLUDED.embedding
        """, (doc_id, source, title, url, content, vector))
        conn.commit()

def search(query, top_k=5):
    qvec = model.encode(query, normalize_embeddings=True).tolist()
    with conn.cursor() as cur:
        cur.execute("""
            SELECT doc_id, title, url, content,
                   1 - (embedding <=> %s::vector) AS score
            FROM rag_chunks
            ORDER BY embedding <=> %s::vector
            LIMIT %s
        """, (qvec, qvec, top_k))
        return cur.fetchall()
